# 文件解析工具（PDF、DOCX）
from pypdf import PdfReader
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from loguru import logger
import zipfile
import xml.etree.ElementTree as ET


class FileParser:
    def parse_pdf(self, file_path: str):
        """
        解析PDF文件

        Args:
            file_path: PDF文件路径

        Returns:
            提取的文本内容
        """
        logger.info(f"开始解析PDF文件：{file_path}")
        reader = PdfReader(file_path)
        pages = reader.pages
        text = ""
        for i, page in enumerate(pages, 1):
            page_text = page.extract_text()
            text += page_text
            logger.info(f"第{i}页提取文本长度：{len(page_text)}")

        logger.info(f"PDF解析完成，总文本长度：{len(text)}")
        return text

    def parse_docx(self, file_path: str):
        """
        解析DOCX文件（包括段落和表格）

        Args:
            file_path: DOCX文件路径

        Returns:
            提取的文本内容
        """
        logger.info(f"开始解析DOCX文件：{file_path}")

        doc = Document(file_path)
        text = ""

        # 提取段落文本
        logger.info(f"文档包含 {len(doc.paragraphs)} 个段落")
        for i, para in enumerate(doc.paragraphs, 1):
            para_text = para.text
            if para_text:
                logger.debug(f"段落 {i}: 长度={len(para_text)}, 内容预览={para_text[:50]}...")
                text += para_text + "\n"
            else:
                logger.debug(f"段落 {i}: 空段落")

        logger.info(f"段落提取完成，文本长度：{len(text)}")

        # 提取表格文本
        logger.info(f"文档包含 {len(doc.tables)} 个表格")
        for i, table in enumerate(doc.tables, 1):
            logger.info(f"正在提取第{i}个表格")
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text
                    if cell_text and cell_text.strip():
                        row_text.append(cell_text.strip())
                if row_text:
                    text += " | ".join(row_text) + "\n"

        # 如果常规方法提取不到文本，尝试从 XML 直接提取
        if len(text.strip()) < 50:  # 如果文本太少，尝试 XML 提取
            logger.info("常规方法提取文本较少，尝试从 XML 直接提取...")
            xml_text = self._extract_text_from_docx_xml(file_path)
            if xml_text and len(xml_text) > len(text):
                text = xml_text
                logger.info(f"从 XML 提取到更多文本，长度：{len(text)}")

        # 尝试提取文档核心属性
        core_props = doc.core_properties
        if core_props.title:
            text += f"\n标题: {core_props.title}\n"
        if core_props.author:
            text += f"作者: {core_props.author}\n"

        logger.info(f"DOCX解析完成，总文本长度：{len(text)}")

        # 如果文本长度为0，记录警告
        if len(text.strip()) == 0:
            logger.warning(f"DOCX文件解析结果为空，可能文件格式不正确或内容为空")
            logger.warning(f"请检查文件是否包含文本内容，或是否为图片/扫描件")

        return text

    def _extract_text_from_docx_xml(self, file_path: str) -> str:
        """
        从 DOCX 文件的 XML 中直接提取所有文本内容
        包括文本框、形状等非常规元素

        Args:
            file_path: DOCX文件路径

        Returns:
            提取的文本内容
        """
        text = ""
        try:
            with zipfile.ZipFile(file_path, 'r') as zf:
                # 读取主文档 XML
                if 'word/document.xml' in zf.namelist():
                    with zf.open('word/document.xml') as f:
                        tree = ET.parse(f)
                        root = tree.getroot()

                        # 提取所有文本节点
                        for elem in root.iter():
                            if elem.tag.endswith('}t'):
                                if elem.text:
                                    text += elem.text
                            elif elem.tag.endswith('}p'):
                                text += "\n"

                # 尝试读取页眉页脚
                for name in zf.namelist():
                    if name.startswith('word/header') or name.startswith('word/footer'):
                        with zf.open(name) as f:
                            tree = ET.parse(f)
                            root = tree.getroot()
                            for elem in root.iter():
                                if elem.tag.endswith('}t'):
                                    if elem.text:
                                        text += elem.text

        except Exception as e:
            logger.error(f"从 XML 提取文本失败: {e}")

        return text
